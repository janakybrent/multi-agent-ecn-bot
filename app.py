# === WRITE STREAMLIT APP ===
app_code = f"""
import streamlit as st
import json
from io import BytesIO
from docx import Document
import openai

st.secrets["OPENAI_API_KEY"]

class AnalystAgent:
    def process(self, ecn_data):
        return {{
            "ecn_id": ecn_data["id"],
            "change_summary": ecn_data["summary"],
            "affected_parts": ", ".join(ecn_data["affected_parts"]),
            "effective_date": ecn_data["effective_date"],
            "owner": ecn_data["owner"]
        }}

class WriterAgent:
    def create_notification(self, data):
        prompt = (
            f"Write a professional customer notification email for the following ECN:\\n"
            f"ECN ID: {{data['ecn_id']}}\\n"
            f"Summary: {{data['change_summary']}}\\n"
            f"Affected Parts: {{data['affected_parts']}}\\n"
            f"Effective Date: {{data['effective_date']}}\\n"
            f"Contact: {{data['owner']}}\\n"
        )
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{{"role": "user", "content": prompt}}],
            max_tokens=500,
            temperature=0.5
        )
        return response["choices"][0]["message"]["content"]

class CRMManagerAgent:
    def create_campaign(self, data):
        return {{
            "campaign_name": f"ECN-{{data['ecn_id']}}-Notification",
            "recipients": ["customer1@example.com", "customer2@example.com"],
            "start_date": data["effective_date"]
        }}

class EmailAgent:
    def send_emails(self, message, recipients):
        return [{{"recipient": r, "status": "sent"}} for r in recipients]

def generate_word_doc(text):
    doc = Document()
    doc.add_heading("Customer Notification", 0)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.set_page_config(page_title="Agentic AI ECN Bot", layout="centered")
st.title("🤖 ECN Notification Bot (Agentic AI)")

with st.expander("📄 ECN JSON Format (click to view example)"):
    st.code('''{{
  "id": "ECN-2048",
  "summary": "Pump housing changed to polycarbonate.",
  "affected_parts": ["P123", "P124"],
  "effective_date": "2025-07-01",
  "owner": "Jane Doe (jane@example.com)"
}}''', language='json')

uploaded_file = st.file_uploader("Upload ECN JSON file", type="json")

if uploaded_file and openai.api_key:
    try:
        ecn_data = json.load(uploaded_file)
        summary = AnalystAgent().process(ecn_data)
        notification = WriterAgent().create_notification(summary)
        campaign = CRMManagerAgent().create_campaign(summary)
        logs = EmailAgent().send_emails(notification, campaign["recipients"])

        st.subheader("📢 Generated Notification")
        st.text_area("Preview", notification, height=300)

        st.subheader("📊 CRM Campaign")
        st.write(f"**Campaign Name:** {{campaign['campaign_name']}}")
        st.write(f"**Start Date:** {{campaign['start_date']}}")
        st.write(f"**Recipients:** {{', '.join(campaign['recipients'])}}")

        st.subheader("📬 Email Log")
        for log in logs:
            st.success(f"✔️ Email sent to {{log['recipient']}}")

        word_buffer = generate_word_doc(notification)
        st.download_button("📥 Download Notification (Word)", word_buffer, file_name="customer_notification.docx")

    except Exception as e:
        st.error(f"❌ Error: {{e}}")
else:
    st.info("Please upload an ECN JSON file and ensure your API key is set.")
"""
with open("app.py", "w") as f:
    f.write(app_code)

# === START STREAMLIT + NGROK ===
from pyngrok import ngrok
import subprocess, time

ngrok.kill()
ngrok.set_auth_token(NGROK_AUTH_TOKEN)
process = subprocess.Popen(["streamlit", "run", "app.py"])
time.sleep(5)
public_url = ngrok.connect(8501)
print("🌐 Your app is live at:", public_url)

